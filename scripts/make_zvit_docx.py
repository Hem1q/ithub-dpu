"""Generate Word .docx from the markdown report using python-docx."""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "docs/zvit_praktyka_oshchadbank.md"
OUT = "docs/Zvit_Praktyka_Oshchadbank.docx"


def set_default_font(doc, name="Times New Roman", size=14):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_page(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def strip_md(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def parse_runs(text):
    """Return list of (text, bold) tuples splitting on **bold** markers."""
    parts = []
    i = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > i:
            parts.append((text[i:m.start()], False))
        parts.append((m.group(1), True))
        i = m.end()
    if i < len(text):
        parts.append((text[i:], False))
    if not parts:
        parts = [(text, False)]
    return [(re.sub(r"`(.+?)`", r"\1", t), b) for t, b in parts]


def add_para(doc, text, *, bold=False, align=None, indent_first=True,
             size=14, caps=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Cm(1.25)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for txt, is_bold in parse_runs(text):
        r = p.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bool(bold or is_bold)
        if caps:
            r.font.all_caps = True
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for a in ("ascii", "hAnsi", "cs", "eastAsia"):
            rfonts.set(qn(f"w:{a}"), "Times New Roman")
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5


def parse_md(md):
    lines = md.split("\n")
    paragraphs = []
    in_code = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            paragraphs.append(("code", line))
            i += 1
            continue
        # markdown table: header | --- | rows
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(
                r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]) and "---" in lines[i + 1]:
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(row)
                j += 1
            paragraphs.append(("table", (header, rows)))
            i = j
            continue
        if not line.strip():
            paragraphs.append(("empty", ""))
        elif line.startswith("# "):
            paragraphs.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            paragraphs.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            paragraphs.append(("h3", line[4:].strip()))
        elif re.match(r"^[-*]\s", line):
            paragraphs.append(("bullet", re.sub(r"^[-*]\s+", "", line)))
        elif re.match(r"^\d+\.\s", line):
            paragraphs.append(("number", re.sub(r"^\d+\.\s+", "", line)))
        elif line.startswith("---"):
            paragraphs.append(("hr", ""))
        else:
            paragraphs.append(("text", line.strip()))
        i += 1
    return paragraphs


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for c, txt in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        for t, b in parse_runs(txt):
            r = p.add_run(t)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(header)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            txt = row[c_idx] if c_idx < len(row) else ""
            for t, b in parse_runs(txt):
                r = p.add_run(t)
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = b


def main():
    with open(SRC, "r", encoding="utf-8") as f:
        md = f.read()

    doc = Document()
    set_default_font(doc)
    set_page(doc)

    paragraphs = parse_md(md)
    last_was_blank = True

    for i, (kind, txt) in enumerate(paragraphs):
        if kind == "empty":
            if not last_was_blank:
                add_blank(doc)
                last_was_blank = True
            continue

        if kind == "h1":
            add_para(doc, txt, bold=True, align="center", indent_first=False,
                     caps=True, space_before=12, space_after=6)
        elif kind == "h2":
            add_para(doc, txt, bold=True, align="left", indent_first=False,
                     space_before=8, space_after=4)
        elif kind == "h3":
            add_para(doc, txt, bold=True, align="left", indent_first=False,
                     space_before=6, space_after=3)
        elif kind == "bullet":
            add_para(doc, "•  " + txt, indent_first=False)
        elif kind == "number":
            add_para(doc, txt, indent_first=True)
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(txt)
            r.font.name = "Courier New"
            r.font.size = Pt(11)
            rpr = r._element.get_or_add_rPr()
            rfonts = OxmlElement("w:rFonts")
            for a in ("ascii", "hAnsi", "cs"):
                rfonts.set(qn(f"w:{a}"), "Courier New")
            rpr.append(rfonts)
        elif kind == "table":
            header, rows = txt
            add_table(doc, header, rows)
        elif kind == "hr":
            add_blank(doc)
        else:
            add_para(doc, txt, indent_first=True)
        last_was_blank = False

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
